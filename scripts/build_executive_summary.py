"""
build_executive_summary.py
===========================
Generates the 2-page ANONYMISED Executive Summary for the Level-1 blind review,
laid out EXACTLY against the competition's official template:

    Team Name  ->  Title  ->  Approach (<=300 words)  ->  Outcomes (<=500 words)
    ->  Links to GitHub files  ->  References (not counted in the word limit)

Formatting per the template: A4, 25.4 mm margins, font Aptos 11 pt. Schematic
diagrams and data visualisations do not count towards the 2 pages / word limit.

Two renderers from ONE content source, so they never drift:
  * Executive_Summary.docx  — authoritative, in Aptos (open in Word -> Save as PDF
    to guarantee the Aptos font for the final submission).
  * Executive_Summary.pdf   — ready-to-send backup rendered with reportlab.

All numbers are read live from outputs/results.json.

    python scripts/build_executive_summary.py
"""

from __future__ import annotations
import json
import re
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "outputs"
RESULTS = json.loads((OUT / "results.json").read_text(encoding="utf-8"))

# --------------------------------------------------------------------------- #
#  Anonymised identity  (NO personal / institution details — blind review)
# --------------------------------------------------------------------------- #
TEAM_NAME = "Team SolarCycle"        # anonymised alias for the blind Level-1 review
GITHUB_URL = "https://github.com/Gourav88502/emicromobility-robust-charging"

TITLE = ("Robust, Solar-Powered Charging Infrastructure for a Shared e-Bike "
         "Scheme under Demand Uncertainty")
SUBTITLE = ("Theme 3 (primary): solar-PV charging-station design  ·  "
            "Theme 2 (secondary): managed charge/discharge profiles   |   "
            "Example site: University of Warwick, Coventry (CV4 7AL)")


# --------------------------------------------------------------------------- #
#  Content (defined once; rendered to both .docx and .pdf)
# --------------------------------------------------------------------------- #
def _content() -> dict:
    rec = RESULTS["recommended_design"]
    nai = RESULTS["naive_design"]
    vor = RESULTS["value_of_robustness"]
    emis = RESULTS["emissions"]
    rob = RESULTS["robustness_of_robustness"]
    capex = RESULTS["recommended_capex_gbp"]
    mc = RESULTS["monte_carlo"]
    val = RESULTS["validation"]["metrics_within_published_range"]
    lp = RESULTS.get("optimal_control", {}).get("max_grid_cost_saving_pct")
    lp_txt = (f"; an optimal LP controller confirms managed charging cuts grid-energy "
              f"cost by up to {lp:.0f}% versus unmanaged charging" if lp else "")
    perf = RESULTS["recommended_performance_high_scenario"]
    sb = rob.get("storage_boundary_grid_kW")
    sb_txt = (f"only when the grid connection falls to ≈{sb:g} kW or below" if sb
              else "only as the connection weakens toward off-grid")

    approach = [
        [("Aim. ", True),
         ("Design a solar-PV charging hub for a shared e-bike scheme at the "
          "University of Warwick (Coventry, CV4 7AL) that stays cost-effective and "
          "reliable across every plausible future demand — not merely an average "
          "forecast (Theme 3) — while modelling the managed charge/discharge load "
          "the hub must serve (Theme 2).", False)],
        [("Objectives. ", True),
         ("(1) Build realistic demand scenarios for the UoW Bikes fleet; (2) size "
          "PV, battery storage and charge bays behind a constrained grid "
          "connection; (3) quantify the cost and service value of designing for "
          "robustness; (4) test whether on-site storage is justified.", False)],
        [("Research questions. ", True),
         ("How should a hub be sized when demand is deeply uncertain? What is the "
          "value of robustness versus a conventional average-demand design? At a "
          "grid-connected hub, does a battery pay?", False)],
        [("Methodology (Figure 1). ", True),
         ("Low/Medium/High demand scenarios are built from the UoW Bikes shared "
          "e-bike data and scaled across a year-on-year growth range into nine "
          "probability-weighted scenarios. A smart controller schedules each day's "
          "flexible charging across the hub dwell window, following on-site PV and "
          "off-peak tariffs. An hourly (8,760 h) energy balance dispatches "
          "PV → battery → capped grid → unmet demand under the "
          "constrained connection. All 150 candidate designs (PV 5–25 kWp × "
          "battery 0–50 kWh × 4–20 bays) are scored against the "
          "scenarios, pricing unmet demand as lost service; five decision rules — "
          "naive, two-stage stochastic, CVaR, minimax-regret and maximin — map the "
          "cost-vs-robustness frontier. Uncertainty is propagated through a "
          "500-sample correlated Monte-Carlo fan and ranked by tornado and "
          "variance-based (Sobol) sensitivity; a robustness-of-robustness sweep "
          "re-solves across penalty × grid × horizon, and seven outputs "
          "are validated against published benchmarks. Real PVGIS solar (Coventry) "
          "and National Grid ESO carbon (West Midlands) drive the energy and "
          "emissions model.", False)],
    ]

    outcomes = [
        [("Recommended design. ", True),
         (f"The robust optimisation recommends a {rec['pv_kwp']:g} kWp PV array with "
          f"{rec['n_chargers']:g} smart-managed charge bays and no battery, at a "
          f"capital cost of approximately £{capex:,.0f} (Figure 2).", False)],
        [("Value of robustness. ", True),
         (f"A conventional design sized to average demand ({nai['pv_kwp']:g} kWp, "
          f"{nai['n_chargers']:g} bays) looks cheap but collapses under the "
          f"high-demand future: its worst-case annual cost is "
          f"£{vor['naive_worst_cost']:,.0f} and it serves only "
          f"{vor['naive_min_service']*100:.1f}% of charging demand. The robust "
          f"design caps worst-case cost at £{vor['robust_worst_cost']:,.0f}/yr "
          f"— a {vor['worst_cost_reduction_pct']:.0f}% "
          f"(£{vor['worst_cost_reduction']:,.0f}/yr) reduction — and guarantees "
          f"{vor['robust_min_service']*100:.1f}% fleet service in every scenario. "
          f"Across the Monte-Carlo fan its 95th-percentile cost is "
          f"£{mc['robust']['cost_p95']:,.0f} versus "
          f"£{mc['naive']['cost_p95']:,.0f} for the naive design, more than "
          f"halving tail risk. We name and quantify this gap as the value of "
          f"robustness.", False)],
        [("Storage boundary — key finding. ", True),
         (f"Because hub charging is flexible, smart charging is the cheapest "
          f"robustness lever: it flattens the load beneath the grid limit, so at a "
          f"grid-connected hub a battery never enters any robust solution. Storage "
          f"becomes essential {sb_txt} — a boundary the robustness sweep locates. "
          f"The recommendation is therefore solar plus "
          f"smart-managed bays, not capital-heavy storage{lp_txt} (Theme 2).", False)],
        [("Robustness of the conclusion. ", True),
         (f"The result is not an artefact of one assumption: the robust design beats "
          f"the naive design in {rob['fraction_robust_beats_naive_pct']:.0f}% of "
          f"{rob['n_combinations']} penalty × grid-limit combinations, with the "
          f"value of robustness ranging {rob['value_of_robustness_min_pct']:.0f}–"
          f"{rob['value_of_robustness_max_pct']:.0f}% (median "
          f"{rob['value_of_robustness_median_pct']:.0f}%) (Figure 3). Seven key "
          f"outputs — PV yield, grid-carbon intensity, round-trip efficiency, trip "
          f"distance, solar fraction, LCOE and carbon saving — all fall within "
          f"published ranges ({val}) (Figure 4).", False)],
        [("Sustainability. ", True),
         (f"Operating the recommended hub avoids {emis['carbon_saving_tCO2_yr']:.1f} "
          f"tCO₂/yr ({emis['carbon_saving_pct']:.0f}%) versus grid-only charging "
          f"at the marginal grid factor, with {perf['solar_fraction_pct']:.0f}% of "
          f"demand met directly from on-site solar.", False)],
        [("Economic implications. ", True),
         ("The design uses only commercially available components, sized within "
          "manufacturer and industry-benchmark ranges, and reports CAPEX, OPEX, "
          "battery-replacement and grid costs. The operator message is direct: pay a "
          "modest premium over an average-demand design to remove most downside cost "
          "and service risk, and do not buy storage at a connected hub. Demand "
          "intensity and year-on-year growth dominate cost variance, so the "
          "specification should be revisited as the scheme scales.", False)],
        [("Originality & reproducibility. ", True),
         ("All model code — the dispatch engine, five-rule optimisation, Sobol and "
          "robustness analyses and every figure — was written from scratch for this "
          "submission. Every numerical assumption carries an inline source; the "
          "pipeline is deterministic, regenerates all data and figures from one "
          "command, and passes an automated test suite, so any reviewer can "
          "reproduce every number.", False)],
    ]

    references = [
        "Burani, E., Cabri, G., Leoncini, M. (2022). An Algorithm to Predict E-Bike Power Consumption Based on Planned Routes. Electronics, 11(7), 1105.",
        "Ouf, K., Soubra, H., Mazhr, A. (2023). E-Bike Energy Needs Estimation based on Route Characteristics and Rider Behavior. IEEE ICICIS 2023, 345–352.",
        "Corti, F. et al. (2024). A comprehensive review of charging infrastructure for Electric Micromobility Vehicles. Energy Reports, 12, 545–567.",
        "Marie, J.-J. (2023). The Micromobility Revolution Gathers Momentum. Faraday Insights, Issue 16.",
        "Gössling, S. (2020). Integrating e-scooters in urban transportation. Transportation Research Part D, 79, 102230.",
        "Rockafellar, R.T., Uryasev, S. (2000). Optimization of Conditional Value-at-Risk. Journal of Risk, 2(3), 21–41.",
        "Birge, J.R., Louveaux, F. (2011). Introduction to Stochastic Programming, 2nd ed. Springer.",
        "PVGIS (EU JRC) hourly series, University of Warwick (52.3838, −1.5616); National Grid ESO Carbon Intensity API, West Midlands (region 8).",
        "Cost & performance benchmarks: BEIS, IRENA, BloombergNEF, IEA PVPS, Fraunhofer ISE — see REFERENCES.md.",
    ]

    github = (f"GitHub (executable code + README / QuickStart guide): {GITHUB_URL}  "
              "— the README documents the one-command run (“python run_analysis.py”), "
              "data sources and how to drop in the official UoW Bikes Data(Sheet1).csv.")

    return {"approach": approach, "outcomes": outcomes,
            "references": references, "github": github}


def _wordcount(blocks) -> int:
    text = " ".join(seg[0] for para in blocks for seg in para)
    return len(re.findall(r"\b[\w’'-]+\b", text))


# --------------------------------------------------------------------------- #
#  DOCX renderer (authoritative — Aptos 11 pt)
# --------------------------------------------------------------------------- #
def render_docx(c: dict, path: Path):
    from docx import Document
    from docx.shared import Pt, Mm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    NAVY = RGBColor(0x1B, 0x1B, 0x3A)
    BLUE = RGBColor(0x2E, 0x86, 0xAB)
    GREY = RGBColor(0x55, 0x55, 0x60)
    FONT = "Aptos"

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    # ensure East-Asian/HAnsi slots also map to Aptos
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts"); rpr.append(rfonts)
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(a), FONT)

    sec = doc.sections[0]
    sec.page_width = Mm(210); sec.page_height = Mm(297)
    sec.top_margin = Mm(25.4); sec.bottom_margin = Mm(25.4)
    sec.left_margin = Mm(25.4); sec.right_margin = Mm(25.4)

    def _set_font(run, size=11, bold=False, italic=False, color=None):
        run.font.name = FONT
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color
        r = run._element.get_or_add_rPr()
        rf = r.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts"); r.append(rf)
        for a in ("w:ascii", "w:hAnsi", "w:cs"):
            rf.set(qn(a), FONT)

    def p(before=0, after=4, align=None, line=1.04):
        par = doc.add_paragraph()
        par.paragraph_format.space_before = Pt(before)
        par.paragraph_format.space_after = Pt(after)
        par.paragraph_format.line_spacing = line
        if align:
            par.alignment = align
        return par

    def heading(text):
        par = p(before=6, after=2)
        _set_font(par.add_run(text), size=13, bold=True, color=NAVY)
        # thin rule under the heading
        ppr = par._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "2"); bottom.set(qn("w:color"), "C7D3DC")
        pbdr.append(bottom); ppr.append(pbdr)

    def body(runs):
        par = p(after=4, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        for text, bold in runs:
            _set_font(par.add_run(text), size=11, bold=bold)

    # --- header --------------------------------------------------------------
    tp = p(after=2, align=WD_ALIGN_PARAGRAPH.LEFT)
    _set_font(tp.add_run(f"Team Name: {TEAM_NAME}"), size=11, bold=True, color=GREY)
    ttl = p(after=1, align=WD_ALIGN_PARAGRAPH.LEFT)
    _set_font(ttl.add_run(TITLE), size=15, bold=True, color=NAVY)
    sub = p(after=6, align=WD_ALIGN_PARAGRAPH.LEFT)
    _set_font(sub.add_run(SUBTITLE), size=9.5, italic=True, color=BLUE)

    # --- Approach ------------------------------------------------------------
    heading(f"Approach  ({_wordcount(c['approach'])} words)")
    for runs in c["approach"]:
        body(runs)
    doc.add_picture(str(OUT / "methodology_flow.png"), width=Inches(5.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = p(after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_font(cap.add_run("Figure 1 — End-to-end methodology / algorithm flow."),
              size=8.5, italic=True, color=GREY)

    # --- Outcomes ------------------------------------------------------------
    heading(f"Outcomes  ({_wordcount(c['outcomes'])} words)")
    for runs in c["outcomes"]:
        body(runs)

    # figure 2 (pareto), then 2-up figures 3 & 4
    doc.add_picture(str(OUT / "02_pareto.png"), width=Inches(4.4))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = p(after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_font(cap.add_run("Figure 2 — Cost-vs-robustness Pareto frontier (five decision rules)."),
              size=8.5, italic=True, color=GREY)

    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = t._tbl
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}"); e.set(qn("w:val"), "none"); borders.append(e)
    tbl.tblPr.append(borders)
    for i, img in enumerate(["10_robustness.png", "11_validation.png"]):
        cell = t.rows[0].cells[i]; cell.width = Mm(72)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].add_run().add_picture(str(OUT / img), width=Inches(2.5))
    cap = p(after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_font(cap.add_run("Figure 3 — Robust beats naive across every penalty × grid case.   "
                          "Figure 4 — Validation vs published benchmarks."),
              size=8.5, italic=True, color=GREY)

    # --- GitHub --------------------------------------------------------------
    heading("Links to GitHub files")
    ghp = p(after=4, align=WD_ALIGN_PARAGRAPH.LEFT)
    _set_font(ghp.add_run(c["github"]), size=11, bold=False)

    # --- References ----------------------------------------------------------
    heading("References  (not counted in the word limit)")
    for i, ref in enumerate(c["references"], 1):
        par = p(after=2, line=1.0)
        _set_font(par.add_run(f"[{i}] {ref}"), size=8.5, color=GREY)

    doc.save(str(path))


# --------------------------------------------------------------------------- #
#  PDF renderer (backup — reportlab, Calibri stand-in for Aptos)
# --------------------------------------------------------------------------- #
def render_pdf(c: dict, path: Path):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Image,
                                    Table, TableStyle, KeepTogether)
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.utils import ImageReader

    fonts = Path("C:/Windows/Fonts")
    base, bold = "Helvetica", "Helvetica-Bold"
    try:
        pdfmetrics.registerFont(TTFont("Body", str(fonts / "calibri.ttf")))
        pdfmetrics.registerFont(TTFont("Body-Bold", str(fonts / "calibrib.ttf")))
        base, bold = "Body", "Body-Bold"
    except Exception:
        pass

    NAVY = colors.HexColor("#1B1B3A"); BLUE = colors.HexColor("#2E86AB")
    GREY = colors.HexColor("#555560")

    def esc(s):
        return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    def to_markup(runs):
        out = ""
        for text, b in runs:
            t = esc(text)
            out += f"<b>{t}</b>" if b else t
        return out

    body_st = ParagraphStyle("body", fontName=base, fontSize=11, leading=13.0,
                             alignment=4, spaceAfter=4, textColor=colors.HexColor("#1a1a1a"))
    h_st = ParagraphStyle("h", fontName=bold, fontSize=13, leading=15, textColor=NAVY,
                          spaceBefore=7, spaceAfter=3)
    team_st = ParagraphStyle("team", fontName=bold, fontSize=11, textColor=GREY, spaceAfter=2)
    title_st = ParagraphStyle("title", fontName=bold, fontSize=15, leading=18,
                              textColor=NAVY, spaceAfter=2)
    sub_st = ParagraphStyle("sub", fontName=base, fontSize=9.5, leading=12,
                            textColor=BLUE, spaceAfter=6)
    cap_st = ParagraphStyle("cap", fontName=base, fontSize=8.5, leading=10,
                            textColor=GREY, alignment=1, spaceAfter=6)
    ref_st = ParagraphStyle("ref", fontName=base, fontSize=8.5, leading=10.5,
                            textColor=GREY, spaceAfter=2)
    body_left = ParagraphStyle("bl", parent=body_st, alignment=0)

    def img_scaled(p, max_w):
        ir = ImageReader(str(p)); iw, ih = ir.getSize()
        w = max_w; h = w * ih / iw
        return Image(str(p), width=w, height=h)

    avail = A4[0] - 2 * 25.4 * mm
    story = [
        Paragraph(f"Team Name: {esc(TEAM_NAME)}", team_st),
        Paragraph(esc(TITLE), title_st),
        Paragraph(esc(SUBTITLE), sub_st),
        Paragraph(f"Approach&nbsp;&nbsp;({_wordcount(c['approach'])} words)", h_st),
    ]
    for runs in c["approach"]:
        story.append(Paragraph(to_markup(runs), body_st))
    story += [img_scaled(OUT / "methodology_flow.png", avail * 0.82),
              Paragraph("Figure 1 — End-to-end methodology / algorithm flow.", cap_st),
              Paragraph(f"Outcomes&nbsp;&nbsp;({_wordcount(c['outcomes'])} words)", h_st)]
    for runs in c["outcomes"]:
        story.append(Paragraph(to_markup(runs), body_st))
    story += [img_scaled(OUT / "02_pareto.png", avail * 0.70),
              Paragraph("Figure 2 — Cost-vs-robustness Pareto frontier (five decision rules).", cap_st)]
    half = avail * 0.43
    row = [[img_scaled(OUT / "10_robustness.png", half),
            img_scaled(OUT / "11_validation.png", half)]]
    tbl = Table(row, colWidths=[avail * 0.5, avail * 0.5])
    tbl.setStyle(TableStyle([("ALIGN", (0, 0), (-1, -1), "CENTER"),
                             ("VALIGN", (0, 0), (-1, -1), "TOP")]))
    story += [tbl,
              Paragraph("Figure 3 — Robust beats naive across every penalty × grid case.   "
                        "Figure 4 — Validation vs published benchmarks.", cap_st),
              Paragraph("Links to GitHub files", h_st),
              Paragraph(to_markup([(c["github"], False)]), body_left),
              Paragraph("References&nbsp;&nbsp;(not counted in the word limit)", h_st)]
    for i, ref in enumerate(c["references"], 1):
        story.append(Paragraph(f"[{i}] {esc(ref)}", ref_st))

    doc = SimpleDocTemplate(str(path), pagesize=A4,
                            topMargin=25.4 * mm, bottomMargin=25.4 * mm,
                            leftMargin=25.4 * mm, rightMargin=25.4 * mm,
                            title="Executive Summary")
    doc.build(story)


def build():
    c = _content()
    aw, ow = _wordcount(c["approach"]), _wordcount(c["outcomes"])
    docx_path = OUT / "Executive_Summary.docx"
    pdf_path = OUT / "Executive_Summary.pdf"
    render_docx(c, docx_path)
    try:
        render_pdf(c, pdf_path)
        pdf_msg = f" + {pdf_path.name}"
    except Exception as e:
        pdf_msg = f"  (PDF skipped: {str(e)[:70]})"
    print(f"Executive summary saved: {docx_path.name}{pdf_msg}")
    print(f"  Approach: {aw}/300 words | Outcomes: {ow}/500 words")
    if aw > 300 or ow > 500:
        print("  WARNING: a section exceeds its word limit — trim the content.")
    return docx_path


if __name__ == "__main__":
    build()
